import io
import docx

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts plain text from DOCX file bytes including table contents.
    """
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    
    # Extract tables which are common in resume formatting
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                # Avoid duplicate adjacent cells due to merged cells
                cleaned_row = []
                for cell_t in row_text:
                    if not cleaned_row or cleaned_row[-1] != cell_t:
                        cleaned_row.append(cell_t)
                text.append(" | ".join(cleaned_row))
                
    return "\n".join(text).strip()
